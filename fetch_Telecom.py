#!/usr/bin/env python
# -*- coding: utf-8 -*-

import re
import datetime
import json
from robobrowser import RoboBrowser
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
import win32com.client

from send_from_gmail import create_message, send_gmail

with open('secret.json') as sf:
    data = json.load(sf)

# 設定項目1：丸三証券のログインID・PWを変数に
koza_id = data['telecom_id']
koza_pw = data['telecom_pw']

# 設定項目2：新聞の種類をタプルに格納
ppr_tpl = ('NKM', 'NSS', 'NRS', 'NKL')
# ppr_tpl=('NKM','NKE','NSS','NRS','NKL','NKP')
# NKM：日本経済新聞朝刊
# NKE：日本経済新聞夕刊
# NSS：日経産業新聞
# NRS：日経ＭＪ（流通新聞）
# NKL：日経地方経済面
# NKP：日経プラスワン


# 設定項目3:何日前の新聞を取得するか
ppr_bfr = 0

# 設定項目4:ファイルの保存フォルダを指定
root = 'C:\\Users\\OWNER\\OneDrive\\Newspapers\\'

# 設定項目5:GmailのID
gmail_id = data['gmail_id']

# 設定項目6:Kindleのメールアドレス ※必ずリストで指定する
kindle_add = data['kindle_addresses']

# 処理開始
# 取得する日の日付を取得
today = datetime.datetime.now()
today -= datetime.timedelta(days=ppr_bfr)
tgt_dt = today.strftime('%Y%m%d')

# ロボブラウザを起動
br = RoboBrowser(parser='lxml', user_agent='a python robot', cache=True, history=False)

# ログイン処理
br.open('https://trade.03trade.com/web/')
form = br.get_form(action='/web/cmnCauSysLgiAction.do')
form['loginTuskKuzNo'].value = koza_id
form['gnziLoginPswd'].value = koza_pw
br.submit_form(form)
print(br.find('title').text)

# 日経テレコンのページへ移動　いくつか遷移する。
br.open('https://trade.03trade.com/web/cmnCauSysLgiSelectAction.do')
tel_url = br.find('a', title='日経テレコン')['onclick'].replace('javascript:window.open(', '') \
    .replace('\'', '').split(',')[0]
br.open(tel_url)

meta_url = br.find_all('meta')[0]['content'].replace('0; url=', '')
br.open(meta_url)

form = br.get_form(action='http://t21.nikkei.co.jp/g3/p03/LCMNDF11.do')
br.submit_form(form)

form = br.get_form(action='LATCA011.do')
br.submit_form(form)

cmp_pprs = []  # 処理済新聞リスト
cnt = 0
for ppr_elm in ppr_tpl:
    cnt += 1
    ppr_url = 'http://t21.nikkei.co.jp/g3/p03/LATCB012.do?mediaCode=' + ppr_elm
    print(ppr_url)

    # 新聞記事のページへ
    br.open(ppr_url)

    # 一番最初の記事から、日付を取得（yyyy/㎜/dd形式）
    info = br.find('li', class_='AttInfoBody').text.replace(u'\xa0', u' ').split(u' ')
    dt = info[0]

    # パンくずリストから新聞名を取得
    ppr = br.find('p', class_='topicPath').find_all('a')[1].text
    print(dt + ppr)

    # 新聞名が取得済リストにないならやる。←無い新聞のページに行ったら朝刊が表示される対策
    if (ppr in cmp_pprs) == False:

        # 日付が取得対象日付だったらやる。違ったらやらない。
        date = datetime.datetime.strptime(dt, '%Y/%m/%d')
        ppr_dt = datetime.date(date.year, date.month, date.day)
        ppr_dt = ppr_dt.strftime('%Y%m%d')

        if tgt_dt == ppr_dt:

            # 面タイトルをリストに格納
            Nav = []
            newsNavs = br.find_all('div', class_='newsNav')
            for newsNav in newsNavs:
                try:
                    newsNav_text = newsNav.find('label').text
                    Nav.append(newsNav_text)
                    print(newsNav_text)
                except:
                    pass

            # 面ごとのタイトルリスト・URLリスト・記事本文リスト・ソースリストを、それぞれリストで格納
            Ttls = []
            Srcs = []
            Txts = []
            newsBlks = br.find_all('ul', class_='listNews valCheck')
            for newsBlk in newsBlks:
                Ttl = []
                Src = []
                Txt = []
                newsIdxs = newsBlk.find_all('li', class_='headlineTwoToneA js-toggle')
                for newsIdx in newsIdxs:
                    try:
                        newsTtl = newsIdx.find('p').find('a').text
                        Ttl.append(newsTtl)
                        newsSrc = newsIdx.find('li', class_='AttInfoBody').text.split(u'\xa0')[1].replace('　', ' ')
                        Src.append(newsSrc)
                        newsUrl = 'http://t21.nikkei.co.jp'+newsIdx.find('a')['href']
                        br.open(newsUrl)
                        print(br.find('h2').text)
                        art_text = br.find('div', class_='col col10 artCSS_Highlight_on').find('p')
                        art_text = str(art_text)  # 注:文字型に変換してあげないとoutertextを扱えない
                        art_text = re.sub(r'</?p?(br)?/?>', '\n', art_text)  # pタグとbrタグを改行に置換
                        Txt.append(art_text)
                    except:
                        pass
                Ttls.append(Ttl)
                Srcs.append(Src)
                Txts.append(Txt)

            if cnt == 1:
                # 繰り返し１回目（最初の新聞）だったら、docxをフォーマットファイルから作成
                print('ファイル更新中...'+dt+ppr)
                doc = Document('format_jp.docx')
                doc.add_page_break()
                doc.add_heading(ppr, level=1)
                doc.add_page_break()
                for i in range(len(Nav)):
                    doc.add_heading(Nav[i], level=2)
                    doc.add_page_break()
                    for j in range(len(Ttls[i])):
                        doc.add_heading(Ttls[i][j], level=3)
                        srcPara = doc.add_paragraph(style='Body Text')
                        srcPara.add_run(dt + ' ' + Srcs[i][j]).font.size = Pt(10)
                        srcPara.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        doc.add_paragraph(Txts[i][j], style='Body Text')
                        doc.add_page_break()

                ppr_name = dt.replace('/', '.') + ' ' + '日本経済新聞' + '.docx'
                path = root + ppr_name
                doc.save(path)

            else:
                # 繰り返し2回目以降だったら、1回目に作ったdocxファイルに追記する。
                # 書き込み処理が冗長なので関数化したいけどまいっか。
                print('ファイル更新中...' +  ppr_name)
                doc = Document(path)
                doc.add_page_break()
                doc.add_heading(ppr, level=1)
                doc.add_page_break()
                for i in range(len(Nav)):
                    doc.add_heading(Nav[i], level=2)
                    doc.add_page_break()
                    for j in range(len(Ttls[i])):
                        doc.add_heading(Ttls[i][j], level=3)
                        srcPara = doc.add_paragraph(style='Body Text')
                        srcPara.add_run(dt + ' ' + Srcs[i][j]).font.size = Pt(10)
                        srcPara.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        doc.add_paragraph(Txts[i][j], style='Body Text')
                        doc.add_page_break()
                doc.save(path)

        else:
            pass

        cmp_pprs.append(ppr)  # 処理済み新聞リストに追加
    else:
        pass

# 目次を更新
print('目次更新中...' + ppr_name)
docApp = win32com.client.Dispatch('Word.Application')
doc = docApp.Documents.Open(path)
doc.TablesOfContents(1).Update()
doc.TablesOfContents(1).UpdatePageNumbers()
doc.Save()
doc.Close()
docApp.Quit()

# Kindleへメールで送る
subject = ppr_name
body = "kindleへ送信"
filename = ppr_name
filepath = path
mine = {'type': 'application', 'subtype': 'vnd.openxmlformats-officedocument.wordprocessingml.document'}
attach_file = {'name': filename, 'path': filepath}
print('メール送信中...' + ppr_name)

msg = create_message(gmail_id, kindle_add, subject, body, mine, attach_file)
send_gmail(gmail_id, kindle_add, msg)

# ログアウト
br.open('https://trade.03trade.com/web/cmnCauSysLgoAction.do')
print(br.find('title').text)
prompt = br.find('h3', class_='function_name').text.encode('shift-jis').decode('shift-jis','replace')
print(prompt)
