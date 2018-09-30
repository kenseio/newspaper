#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
import re
import time
import datetime
import json
from robobrowser import RoboBrowser
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import win32com.client

from send_from_gmail import create_message, send_gmail
from pil_for_kindle import image_process

with open('secret.json') as sf:
    data = json.load(sf)

root = 'C:\\Users\\OWNER\\OneDrive\\Newspapers\\'

# 設定項目1:見出しページをリストに格納
lstSectionURL = ['/news/national/', '/news/business/', '/news/world/', '/news/asia-pacific/', '/opinion/', '/life/']

# 設定項目5:GmailのID
gmail_id = data['gmail_id']

# 設定項目6:Kindleのメールアドレス ※必ずリストで指定する
kindle_add = data['kindle_addresses']

# メインのブラウザ
br = RoboBrowser(parser='lxml', user_agent='a python robot', cache=True, history=True)
strRoot = 'http://www.japantimes.co.jp'  # 固定

# 画像を一時的に読み込むブラウザ、保存するときのファイル名
img_path = root + 'img.jpg'

# 前回実行日時を読み込む
fileLastDate = open('LastSubmitDate_JapanTimes.txt', 'r')
dtLastDate = fileLastDate.read()
fileLastDate.close()
print("/---" + str(dtLastDate) + "以降に更新された記事を読み込みます")

# docxフォーマットを開く
doc = Document('format_en.docx')

# ---- 見出し毎の繰り返しここから

for strSectionUrl in lstSectionURL:
    print('/---' + strSectionUrl.replace('/', ' ') + 'を実行します')
    # 記事のURLを取得しリストに格納
    # OpinionとLifeは1ページだけ。他は1ページ目と2ページ目を読み込む。
    lstArticleUrl = []
    if strSectionUrl in ['/opinion/', '/life/']:
        br.open(strRoot + strSectionUrl)
        strCntnt = br.find('h1', class_='page-title').text
        strCntnt = re.sub(r"\n|  ", '', strCntnt).strip()
        objSect = br.find('div', id='wrapper')
        for objHgroup in objSect.find_all('hgroup'):
            lstArticleUrl.append(objHgroup.find('p').find('a')['href'])

    else:
        for i in range(2):
            br.open(strRoot + strSectionUrl + 'page/' + str(i+1) + '/')
            strCntnt = br.find('h1', class_='page-title').text
            strCntnt = re.sub(r"\n|  ", '', strCntnt).strip()
            objSect = br.find('section')
            for objHgroup in objSect.find_all('hgroup'):
                lstArticleUrl.append(objHgroup.find('p').find('a')['href'])

    # 見出し１を作る
    doc.add_page_break()
    doc.add_heading(strCntnt, level=1)
    doc.add_page_break()

    # リストに格納した記事のURLにアクセス
    # 記事の日付を最初に見て、前回実行日時以降のものだったら実行。
    # 前回実行日時以前だったらループを中断
    for strArtcleUrl in lstArticleUrl:
        try:
            br.open(strArtcleUrl)
        except:
            print("/---記事読み込み失敗 スキップします")
            continue

        objArticle = br.find('article', role='main')  # 記事のタイトル・画像・記者など
        objBody = br.find('div', id='jtarticle')  # 記事本文
        try:
            dtArticleDate = objArticle.find('time')['datetime'].replace('T', ' ')[0:19]
            # dtArticleDate=datetime.datetime.strptime(strArticleDate, '%Y-%m-%d %H:%M:%S')

        except:
            dtArticleDate = dtLastDate  # 日付取得でエラーになったら、前回実行日時を指定してとりあえず読み込ませる

        print("/---記事日付：" + dtArticleDate)

        # タイトル
        try:
            strTitle = objArticle.find('h1').text
            print("/---タイトル：" + strTitle.replace('\xa5', '\\').replace('\u2014', '-').replace('\u20ac', 'euro'))
        except:
            print("/---記事読み込み失敗 スキップします")
            continue

        if dtArticleDate >= dtLastDate:  # 記事日時と前回実行日時を比較。ここで判定。
            print("/---読み込みます")
            # タイトル
            doc.add_heading(strTitle, level=2)

            # 記事日時と見出し
            dtDate = datetime.datetime.strptime(dtArticleDate[0:19], '%Y-%m-%d %H:%M:%S')
            strDate = datetime.datetime.strftime(dtDate, '%B %d, %Y  %H:%M:%S')

            paraDate = doc.add_paragraph(style='Body Text')
            paraDate.add_run(strDate + ' | ' + strCntnt ).font.size = Pt(9)
            paraDate.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # 画像
            try:
                strImgSrc = objArticle.find('figure').find('img')['src']
                image_process(root, strImgSrc)
                doc.add_picture(img_path)
                time.sleep(0.2)
                os.remove(img_path)
            except:
                pass

            # 画像説明文
            try:
                strImgCpt = objArticle.find('figure').find('figcaption').text
                paraImgCpt = doc.add_paragraph(style='Body Text')
                paraImgCpt.add_run(strImgCpt).font.size = Pt(8)
            except:
                pass

            # 記事本文
            for objElm in objBody.find_all('p'):
                strAtclText = objElm.text.strip()
                paraAtclTxt = doc.add_paragraph(style='Body Text')
                paraAtclTxt.add_run(strAtclText).font.size = Pt(11)

            # クレジットとライター
            try:
                strCredit = objArticle.find('p', class_='credit').text
            except:
                strCredit = ''
            try:
                strWriter = objArticle.find('h5').text
            except:
                strWriter = ''
            if len(strCredit + strWriter) > 0:
                paraAtclTxt = doc.add_paragraph(style='Body Text')
                paraAtclTxt.add_run('(' + strCredit + ' ' + strWriter + ')').font.size = Pt(10)
            doc.add_page_break()

        else:
            print("/---処理をスキップします")
            continue


# ---- 見出し毎の繰り返しここまで

# docxファイルを保存
today = datetime.datetime.today()
strToday = today.strftime('%Y.%m.%d')

ppr_name = strToday+' JapanTimes.docx'
path = root + ppr_name
doc.save(path)

# 目次を更新
print("/---目次を更新中：" + ppr_name)
docApp = win32com.client.Dispatch('Word.Application')
doc = docApp.Documents.Open(path)
doc.TablesOfContents(1).Update()
doc.TablesOfContents(1).UpdatePageNumbers()
doc.Save()
doc.Close()
docApp.Quit()

# Kindleへメールで送る
subject = ppr_name.replace('.docx', '')
body = "kindleへ送信"
filename = ppr_name
filepath = path
mine = {'type': 'application', 'subtype': 'vnd.openxmlformats-officedocument.wordprocessingml.document'}
attach_file = {'name': filename, 'path': filepath}
print('/---メール送信中：' + ppr_name)

msg = create_message(gmail_id, kindle_add, subject, body, mine, attach_file)
send_gmail(gmail_id, kindle_add, msg)

# 今回実行日時をファイルに書き込む
fileLastDate = open('LastSubmitDate_JapanTimes.txt', 'w')
dtLastDate = fileLastDate.write(str(datetime.datetime.now()))
fileLastDate.close()
print('/---今回実行日時は：' + str(datetime.datetime.now()))

print('/---処理を終了しました。')
