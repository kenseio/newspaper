#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
import re
import time
import datetime
import json
from pytz import timezone
from robobrowser import RoboBrowser
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import win32com.client

from send_from_gmail import create_message, send_gmail
from pil_for_kindle import image_process

with open('secret.json') as sf:
    data = json.load(sf)

root = 'K:\\新聞\\AsiaNikkei\\'

# 設定項目1:見出しページをリストに格納
lstSectionURL = ['/Editor-s-Picks/', '/Business/', '/Economy/', '/Politics/']

# 設定項目5:GmailのID
gmail_id = data['gmail_id']

# 設定項目6:Kindleのメールアドレス ※必ずリストで指定する
kindle_add = data['kindle_addresses']

# メインのブラウザ
br = RoboBrowser(parser='lxml', user_agent='a python robot', cache=True, history=True)
strRoot = 'http://asia.nikkei.com'  # 固定

# 画像を一時的に保存するときのファイル名
img_path = root + 'img.jpg'

# 前回実行日時を読み込む
fileLastDate = open('LastSubmitDate_AsiaNikkei.txt', 'r')
strLastDate = fileLastDate.read()
dtLastDate = datetime.datetime.strptime(strLastDate[0:19], '%Y-%m-%d %H:%M:%S').astimezone(timezone('Asia/Tokyo'))
fileLastDate.close()
print("/---" + str(strLastDate) + "以降に更新された記事を読み込みます")

# docxフォーマットを開く
doc = Document('format_en.docx')

# ---- 見出し毎の繰り返しここから
for SectionURL in lstSectionURL:
    print('/---' + SectionURL.replace('/', ' ') + 'を実行します')
    # 見出しページの2ページ目まで読み込んで、記事の日付とURLを取得し、それぞれリストに格納
    lstTitle = []
    lstArticleUrl = []
    lstDt = []
    for i in range(2):
        strSect = SectionURL + '?page=' + str(i + 1)
        br.open(strRoot + strSect)
        strCntnt = br.find('header', class_='content__header').find('span', class_='ezstring-field').text
        objCntnt = br.find('section', id='article-stream')
        # aタグかつ、title属性があって、テキストもあるものを全て拾う。写真のリンクは拾いたくない。
        tagTitles = objCntnt.find_all('a', title=re.compile('\w*'), text=re.compile('\w*'))
        for tagTitle in tagTitles:
            if tagTitle.parent.name != 'li':
                lstTitle.append(tagTitle.text)
                lstArticleUrl.append(tagTitle.parent.find('a')['href'])
                strArticleDate = tagTitle.find_next('time')['data-time-utc']
                dtArticleDate = (datetime.datetime.strptime(strArticleDate, '%B %d, %Y %H:%M %Z')
                                 + datetime.timedelta(hours=9)).astimezone(timezone('Asia/Tokyo'))
                lstDt.append(dtArticleDate)

    # 見出し1を作る
    doc.add_page_break()
    doc.add_heading(strCntnt, level=1)
    doc.add_page_break()

    # 記事日付が前回実行日時より新しいものを読み込んで、docxに書き出し。
    for j in range(len(lstDt)):
        print("/---" + str(j + 1) + "個目の記事を処理します")
        print("/---記事日付：" + str(lstDt[j]))
        print("/---タイトル：" + lstTitle[j])

        print(lstDt[j], dtLastDate)
        if lstDt[j] > dtLastDate:
            print("/---読み込みます")
            print("/---記事URL：" + strRoot + lstArticleUrl[j])

            # 記事URLを開いて、NOT FOUNDだったらスキップ
            br.open(strRoot + lstArticleUrl[j])
            if '200' not in str(br.response):
                print("/---記事URLがNot Foundでした")
                continue

            # タイトル
            strTitle = br.find('h1', class_='article__title').text
            doc.add_heading(strTitle, level=2)

            # 日付+見出し
            strTopic = br.find('span', class_='article__topic').text.strip()
            strDate = br.find('div', class_='article__details').find('time').text
            strDate = re.sub(r"\n|  ", '', strDate).strip()
            paraDate = doc.add_paragraph(style='Body Text')
            paraDate.add_run(strDate + ' | ' + strTopic).font.size = Pt(9)
            paraDate.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # サブタイトル
            try:
                strSubtitle = br.find('p', class_='article__sub-title').text
                paraSubtitle = doc.add_paragraph(style='Body Text')
                paraSubtitle.add_run(strSubtitle).font.size = Pt(12)
            except:
                print('/---Note:サブタイトル無かった')

            # Articleタグ配下で最初に出てくる画像を探す。
            # 本文の画像は別途取得するので、親タグのクラスが"article" "article__content"のどっちか。
            # ただし、無いかもしれない。
            try:
                objImg = br.find('div', class_='article').find('img')
                if objImg.parent['class'][0] in ['article', 'article__content']:
                    strImgSrc = objImg['src']
                    image_process(root, strImgSrc)
                    doc.add_picture(img_path)
                    time.sleep(0.2)
                    os.remove(img_path)
                else:
                    print('/---Note:画像無かったよ')
            except:
                print('/---Note:画像無かったよ')

            # 画像説明文
            try:
                objImgCpt = br.find('div', class_='article').find('span', class_='article__caption')
                if objImgCpt.parent['class'][0] in ['article', 'article__content']:
                    strImgCpt = objImgCpt.text
                    strImgCpt = re.sub(r"\t|\n|\xa0|\xa9|  ", '', strImgCpt).strip()
                    paraImgCpt = doc.add_paragraph(style='Body Text')
                    paraImgCpt.add_run(strImgCpt).font.size = Pt(8)
                else:
                    print('/---Note:画像説明文無かったよ')
            except:
                print('/---Note:画像説明文無かったよ')

            # 本文中のテキスト・画像・画像説明文
            objArticle = br.find('div', class_='ezrichtext-field')
            for objElm in objArticle.descendants:
                if objElm.name == 'p':
                    strArticleText = objElm.text
                    paraArticleText = doc.add_paragraph(style='Body Text')
                    paraArticleText.add_run(strArticleText).font.size = Pt(11)

                elif objElm.name == 'img':
                    strImgSrc = objElm['src']
                    image_process(root, strImgSrc)
                    doc.add_picture(img_path)
                    time.sleep(0.2)
                    os.remove(img_path)

                elif objElm.name == 'span' and objElm['class'][0] == 'article__caption':
                    strImgCpt = objElm.text
                    strImgCpt = re.sub(r"\t|\n|\xa0|\xa9|  ", '', strImgCpt).strip()
                    paraImgCpt = doc.add_paragraph(style='Body Text')
                    paraImgCpt.add_run(strImgCpt).font.size = Pt(8)

            # author
            try:
                strAuthor = br.find('div', class_='article__details').find('span', class_='article__author').text.strip()
                paraArticleText = doc.add_paragraph(style='Body Text')
                paraArticleText.add_run('(' + strAuthor + ')').font.size = Pt(10)
            except:
                print('/---Note:Author無かった')

            doc.add_page_break()

        #ひとつの記事ここまで

        else:
            print("/---読み込みません")

# ---- 見出し毎の繰り返しここまで

today = datetime.datetime.today()
strToday = today.strftime('%Y.%m.%d')

ppr_name = strToday + ' NikkeiAsianReview.docx'
path = root + ppr_name
doc.save(path)

# 目次を更新
print("/---目次を更新中：" + ppr_name)
docApp = win32com.client.Dispatch('Word.Application')
doc = docApp.Documents.Open(path)
doc.TablesOfContents(1).Update()
doc.TablesOfContents(1).UpdatePageNumbers()
doc.Save()
doc.Close()
docApp.Quit()

# Kindleへメールで送る
subject = ppr_name.replace('.docx', '')
body = "kindleへ送信"
filename = ppr_name
filepath = path
mine = {'type': 'application', 'subtype': 'vnd.openxmlformats-officedocument.wordprocessingml.document'}
attach_file = {'name': filename, 'path': filepath}
print('/---メール送信中：' + ppr_name)

msg = create_message(gmail_id, kindle_add, subject, body, mine, attach_file)
send_gmail(gmail_id, kindle_add, msg)

# 今回実行日時をファイルに書き込む
fileLastDate = open('LastSubmitDate_AsiaNikkei.txt', 'w')
dtLastDate = fileLastDate.write(str(datetime.datetime.now()))
fileLastDate.close()
print('/---今回実行日時は：' + str(datetime.datetime.now()))

print('/---処理を終了しました。')
